##현재 해야 하는 테스트
#
#이미지 자동조정 테스트 (너무크면 줄이고, 디시콘인것도 조정할것)
#docx 하이퍼링크 테스트
#
#넣어야 하는데 귀찮은 것: 줄긋기, 색깔글씨
#
# 글의 댓글을 긁어와 docx에 작성하는 테스트
#
# 보이스 리플 인식시키기 및 보이스리플입니다로 표시만 하기 _해결 완료
#
# 아 g를 지역변수명으로 씀...
# 예외처리  \ / : * ? " < > |


import requests
from bs4 import BeautifulSoup
from selenium import webdriver
import json
import os
import shutil
from time import sleep
from PIL import Image
import copy

from docx import Document
import docx
from docx.oxml.ns import qn, nsdecls

#임시로 해놓은 import들 안쓰는 것 있으면 삭제할 것!
print("세팅중..")

MAIN_URL = "https://m.dcinside.com/board/jumper/번호" #상관없음

#MAIN_URL = str(input('디시글 URL을 입력하시오..:'))

#생성할 docx 위치 설정
COMMENT_DCCON_SIZE = 40
DOCX_SAVE_PATH = r"docx 파일 저장 위치"
docx_name = "testing"
DRIVER_PATH = r'크롬 드라이버 위치'
IMG_TEMP_PATH = r'다운로드할 사진들을 임시로 저장하는 폴더'

wait_Time = 4 #대기시간(4~6초 권장)

file_path_list = [r'docx 파일들이 저장된 곳의 주소',]

#docx 생성 테스트
#굴림체, 바탕체, 돋움체, 궁서체, Arial, Verdana, Courier New

galldocx = Document()

DEFAULT_TEXT_STYLE = { #기본 텍스트 스타일 
	'font':'Arial', #영어 글꼴
	'hfont':'굴림체', #한글 글꼴 
	'size':10 #글씨 크기

}
DEFAULT_TITLE_TEXT_STYLE = { #기본 텍스트 스타일 
	'font':'Arial', #영어 글꼴
	'hfont':'돋움체', #한글 글꼴 
	'size':22 #글씨 크기
}
DEFAULT_COMMENT_TEXT_STYLE = { #기본 텍스트 스타일 
	'font':'Arial', #영어 글꼴
	'hfont':'돋움체', #한글 글꼴 
	'size':10, #글씨 크기
	'bold':True
}

#파일이름중 쓰면 안되는 특수기호를 적당히 바꾸어 줍니다.
def changeFileName(g_filename):
	for (char_in,char_out) in [('\\','_'),('/','_'),(':',';'),('*','_'),('"',"'"),('<','{'),('>','}'),('|',"l"),('?',"_")]:
		if char_in in g_filename:
			g_filename=g_filename.replace(char_in,char_out)
	return g_filename

#함수
def setText(setrun, textstyle = {}):
	'''run 에 들어가는 텍스트의 속성을 바꾸어 줍니다.
	textstyle는 딕셔너리 자료형이 들어가야 합니다.'''

	#영어 글꼴 바꾸기
	if 'font' in textstyle: 
		setrun.font.name = textstyle['font']
	else:
		setrun.font.name = DEFAULT_TEXT_STYLE['font']
	
	#한글 글꼴 바꾸기
	if 'hfont' in textstyle: 
		setrun.element.rPr.rFonts.set(qn('w:eastAsia'), textstyle['hfont'])
	else:
		setrun.element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_TEXT_STYLE['hfont'])

	#글씨 크기 바꾸기
	if 'size' in textstyle:
		setrun.font.size = docx.shared.Pt(textstyle['size'])
	else:
		setrun.font.size = docx.shared.Pt(DEFAULT_TEXT_STYLE['size'])

	#bold 기본값-false
	if 'bold' in textstyle:
		setrun.bold = textstyle['bold']
	else:
		setrun.bold = False

	#italic 기본값-False
	if 'italic' in textstyle:
		setrun.italic = textstyle['italic']
	else:
		setrun.italic = False

	#color
	if 'color' in textstyle:
		setrun.font.color.rgb = docx.shared.RGBColor(textstyle['color'][0],textstyle['color'][1],textstyle['color'][2])

#폰트 확인 출력값은 리스트(각 원소는 튜플로('속성',값)으로 존재)
def findFont(g_str):
	g_output = []

	#글씨체 확인  폰트달라지는거 추가해야 함!
	if 'Gulim' in g_str:
		g_output.append(('hfont','굴림체'))
	elif 'Dotum' in g_str:
		g_output.append(('hfont','돋움체'))
	elif 'Gungsuh' in g_str:
		g_output.append(('hfont','궁서체'))
	elif 'Dotum' in g_str:
		g_output.append(('hfont','돋움체'))
	elif 'Arial' in g_str:
		g_output.append(('font','Arial'))
	elif 'Verdana' in g_str:
		g_output.append(('font','Verdana'))
	elif 'Courier New' in g_str:
		g_output.append(('font','Courier New'))

	#글자 크기 확인
	if 'font-size' in g_str:
		if 'px' in g_str:
			g_fontsize = g_str[g_str.find('font-size')+10:g_str.find('px')]
		if 'pt' in g_str:
			g_fontsize = g_str[g_str.find('font-size')+10:g_str.find('pt')]
		g_output.append(('size',int(g_fontsize)))

	#색깔 배경, 글자 확인
	if 'color' in g_str:
		pass

	return g_output

#다운로드가 다 될때까지 기다리고, 다운로드 한 파일 정보를 반환합니다.
def waitFileDownload(g_driver,URL,set_filename_URL=True,sleep_time_cycle=0.2):
	#파일 정보는 디렉토리 형태로 name, size, file_path, format 이 보내집니다.
	g_file_info = {}
	global IMG_TEMP_PATH

	print("다운로드 하기 전 현재 탭 개수:",len(g_driver.window_handles))

	#기존 디렉토리에 있는 파일 확인
	g_check_file_list = os.listdir(IMG_TEMP_PATH)
	g_check_crdownload = []
	#.crdownload파일이 다운받기 전에 존재하는 지 확인, 있으면 체크하기
	for g_temp_filename in g_check_file_list:
		if 'crdownload' == os.path.splitext(g_temp_filename)[-1]:
			g_check_crdownload.append(g_temp_filename)

	#dcimg가 막히는 경우가 있어서 우회하는 방법 임시방편으로 사용하기..
	URL=URL.replace("dcimg4","dcimg3")
	URL=URL.replace("dcimg3","dcimg7")
	#옛날 디시콘을 불러올시 에러가 발생하는 것 같음.. 잠시 dcimg2에러 발생 
	if 'dcimg2' in URL: _=0/0
	URL=URL.replace("dcimg6","dcimg7")
	URL=URL.replace("dcimg7","dcimg4")
	g_driver.get(URL)
	g_download_check = True

	#셀레니움에서 새탭을 열고 이미지를 다운받는 과정
	g_driver.execute_script('window.open("about:blank");')
	g_driver.switch_to.window(g_driver.window_handles[1])
	
	#IMG_TEMP_PATH
	print('다운로드중..\n다운로드 완료시에도 넘어가지 않는다면 crdownload 파일이 있는지 확인해 보시기 바랍니다.')
	while g_download_check:
		g_download_check = False
		sleep(sleep_time_cycle)
		g_current_file_list = os.listdir(IMG_TEMP_PATH)
		for g_temp_filename in g_current_file_list:
			if g_temp_filename not in g_check_file_list:
				print("현재 파일 상태: ",os.path.splitext(g_temp_filename)[-1])
				if '.crdownload' == os.path.splitext(g_temp_filename)[-1]:
					g_download_check = True #아직 다운로드가 완료되지 않음
				if '.tmp' == os.path.splitext(g_temp_filename)[-1]:
					g_download_check = True #아직 다운로드가 완료되지 않음
		if len(g_check_file_list) == len(g_current_file_list):
			g_download_check = True

	print('다운로드 완료. 더 이상 넘어가지 않는다면')
	#다운로드한 파일을 찾은후 파일 속성 기록하기
	#만약에 바뀐 파일이 한개 이상이면 예외처리? 그냥 처음나온 파일만 가져다 쓰기?
	while len(g_check_file_list) == len(g_current_file_list):
		sleep(sleep_time_cycle)
		g_current_file_list = os.listdir(IMG_TEMP_PATH)
	g_current_file_list = os.listdir(IMG_TEMP_PATH)
	g_check_newdownload = []
	g_newdownload_path = ''
	for g_temp_filename in g_current_file_list:
		if g_temp_filename not in g_check_file_list:
			g_check_newdownload.append(g_temp_filename)

	if len(g_check_newdownload) == 1: #매우 편ㅡ안하게 처리됨
		g_newdownload_path = IMG_TEMP_PATH + '\\' + g_check_newdownload[0] 
	elif len(g_check_newdownload) == 0:
		print("Error: 다운로드 한 파일이 없거나 사라졌습니다.")
		g_driver.close()
		g_driver.switch_to.window(g_driver.window_handles[0])
		return {}
	else:
		print("Warning: 두개 이상의 파일이 작동중에 다운로드 되었습니다.")
		g_newdownload_path = IMG_TEMP_PATH + '\\' + g_check_newdownload[0] 

	#다운로드 완료, 탭 닫기

	print("현재 탭 개수:",len(g_driver.window_handles))
	if len(g_driver.window_handles) != 1:
		while len(g_driver.window_handles) > 1:
			g_driver.switch_to.window(g_driver.window_handles[-1])
			g_driver.close()
	g_driver.switch_to.window(g_driver.window_handles[0])

	print(g_check_newdownload)
	print(os.path.splitext(g_newdownload_path))
	print(g_newdownload_path)
	g_file_info['name'], g_file_info['format'] = os.path.splitext(g_newdownload_path)

	#파일명을 바꾸는 과정
	if set_filename_URL:
		os.rename(g_newdownload_path,IMG_TEMP_PATH + '\\' + URL[-20:] + g_file_info['format'])
		g_newdownload_path = IMG_TEMP_PATH + '\\' + URL[-20:] + g_file_info['format']
		g_file_info['name'] = URL[-20:]
	g_file_info['file_path'] = g_newdownload_path
	g_image = Image.open(g_file_info['file_path'])
	g_file_info['size'] = g_image.size
	g_image.close()
	return g_file_info

#사이즈 자동 조정(예정)
image_in_comment = False
def resizeImage(g_size):
	g_temp_size = [g_size[0],g_size[1]]
	if image_in_comment:
		g_temp_size = [COMMENT_DCCON_SIZE,COMMENT_DCCON_SIZE]
	else:
		g_temp_size[0]/=2
		g_temp_size[1]/=2

	if g_temp_size[0]>=450:
		print("사진이 너무 큽니다.. 사진 사이즈를 조정합니다.")
		g_temp_size[1] = (g_temp_size[1]*450)/g_temp_size[0]
		g_temp_size[0] = 450

	return g_temp_size

#글 또는 대슬 작성자에 대한 정보를 딕셔너리 형태로 추출합니다
def getIdInfo(g_html,g_is_comment = False):
	g_temp = g_html.find(attrs={'class':'gall_writer ub-writer'})
	if g_temp == None:
		g_temp = g_html.find(attrs={'class':'gall_writer ub-writer hitnik fix'})
	g_in_info = {}
	g_in_info['user-nickname'] = g_temp['data-nick']
	g_in_info['userid'] = g_temp['data-uid']
	g_in_info['userip'] = g_temp['data-ip']
	g_in_info['user'] = g_in_info['userid'] if not g_in_info['userip'] else g_in_info['userip']
	g_in_info['date'] = ' '
	if not g_is_comment :
		g_temp = g_html.find(attrs={'class':'gall_date'})
		g_in_info['date'] = g_temp['title']
	else:
		g_temp = g_html.find(attrs={'class':'fr clear'})
		g_in_info['date'] = g_temp.get_text()[:-2]
	return g_in_info

#재귀함수를 이용한 태그 분석
def listingTag(g_paragraph,g_temp_tag,g_TEXT_STYLE=copy.deepcopy(DEFAULT_TEXT_STYLE)):
	global paragraph, run, driver, chrome_options, settings, prefs, IMG_TEMP_PATH, wait_Time
	
	#크로미움이 갑자기 닫히는 경우가 발생함
	try:
		_=len(driver.window_handles)
	except:
		print("예상치 못하게 크로미움이 종료되었습니다.. 다시 여는중")
		driver = webdriver.Chrome(chrome_options=chrome_options,executable_path=DRIVER_PATH)
		driver.implicitly_wait(wait_Time)
		driver.set_window_size(300, 100) 

	g_Tag_type = str(g_temp_tag.name) #현재 최상위 태그를 확인합니다.
	g_g_TEXT_STYLE = copy.deepcopy(g_TEXT_STYLE) #태그로 인해 글꼴이 바뀔때 저장하는 것

	print("타입 확인:",g_Tag_type)

	if g_Tag_type == 'None': #그냥 글자일 경우
		run = g_paragraph.add_run(g_temp_tag.string)
		setText(run,g_TEXT_STYLE)

	elif g_Tag_type =='a':
		#속성 확인 해야함!
		try:
			run = g_paragraph.add_run('hyperlink: '+g_temp_tag['href']+'\n')
		except:
			pass
		else:
			setText(run,{'bold':True})
		#속성 확인 종료
		for g_tag_in_tag in list(g_temp_tag):
			listingTag(g_paragraph,g_tag_in_tag,g_TEXT_STYLE)


	elif g_Tag_type =='p':
		#속성 확인 해야함!
		#속성 확인 종료
		for g_tag_in_tag in list(g_temp_tag):
			listingTag(g_paragraph,g_tag_in_tag,g_TEXT_STYLE)

	elif g_Tag_type =='div':
		#속성 확인 해야함!
		#속성 확인 종료
		run = g_paragraph.add_run('\n')
		for g_tag_in_tag in list(g_temp_tag):
			listingTag(g_paragraph,g_tag_in_tag,g_TEXT_STYLE)

	elif g_Tag_type =='span':
		#속성 확인 해야함!
		try:
			g_font = findFont(str(g_temp_tag['style']))
		except:
			print("error: There are no 'style' in style tag")
		else:
			#얻은 결과를 적용합니다.
			for g_font_i in g_font:
				g_g_TEXT_STYLE[g_font_i[0]] = g_font_i[1]

		#속성 확인 종료
		#print("span 내의 문자열 확인: " g_temp_tag.string)
		for g_tag_in_tag in list(g_temp_tag):
			listingTag(g_paragraph,g_tag_in_tag,copy.deepcopy(g_g_TEXT_STYLE))

	elif g_Tag_type == 'font': #아직은 넘어가기
		#속성 확인 (글씨체만)
		try:
			g_font = findFont(str(g_temp_tag['face']))
		except:
			print("error: There are no 'face' in font tag")
		else:
			#얻은 결과를 적용합니다.
			for g_font_i in g_font:
				g_g_TEXT_STYLE[g_font_i[0]] = g_font_i[1]

		#속성 확인 종료
		for g_tag_in_tag in list(g_temp_tag):
			listingTag(g_paragraph,g_tag_in_tag,copy.deepcopy(g_g_TEXT_STYLE))

	elif g_Tag_type == 'br': #줄바꿈? 아직 분석중
		#어떻게 해야할지 고민하기!!
		run = g_paragraph.add_run('\n')

	#bold
	elif g_Tag_type == 'b':
		#더 나은 방법 없나..
		g_g_TEXT_STYLE['bold'] = True
		for g_tag_in_tag in list(g_temp_tag):
			listingTag(g_paragraph,g_tag_in_tag,copy.deepcopy(g_g_TEXT_STYLE))
		del(g_g_TEXT_STYLE)

	#italic
	elif g_Tag_type == 'i':
		#더 나은 방법 없나..
		g_g_TEXT_STYLE['italic'] = True
		for g_tag_in_tag in list(g_temp_tag):
			listingTag(g_paragraph,g_tag_in_tag,copy.deepcopy(g_g_TEXT_STYLE))


	elif g_Tag_type =='img':
		print('img통과')
		#디시콘인지 확인
		g_img_tag = ''
		try:
			g_img_tag = g_temp_tag['class']
		except:
			print('@@@@@@@@@@@@@사진 추출중@@@@@@@@@@@@')
			print(g_temp_tag['src'])
			g_img_tag = g_temp_tag['src']
			#새탭 열기
			try:
				g_docximage_info = waitFileDownload(driver,g_img_tag)
				print(g_docximage_info['size'])
				g_docximage_info['size'] = resizeImage(g_docximage_info['size'])
				print(g_docximage_info)
				galldocx.add_picture(g_docximage_info['file_path'],width= docx.shared.Pt(g_docximage_info['size'][0]), height= docx.shared.Pt(g_docximage_info['size'][1]))
				paragraph = galldocx.add_paragraph()
				os.remove(g_docximage_info['file_path'])
			except:
				print('###############이미지를 불러오는 중에 에러가 발생했습니다.')
				sleep(1)

		else:
			print("@@@@@@@@@@@@@디시콘 다운중@@@@@@@@@@@")
			print(g_temp_tag['src'])
			g_img_tag = g_temp_tag['src']
			#새탭 열기
			try:
				g_docximage_info = waitFileDownload(driver,g_img_tag)
				g_docximage_info['size'] = resizeImage(g_docximage_info['size'])
				print(g_docximage_info)
				galldocx.add_picture(g_docximage_info['file_path'],width= docx.shared.Pt(g_docximage_info['size'][0]), height= docx.shared.Pt(g_docximage_info['size'][1]))
				paragraph = galldocx.add_paragraph()
				os.remove(g_docximage_info['file_path'])
			except:
				print('###############디시콘을 불러오는 중에 에러가 발생했습니다.')
				print('###############디시콘이 소멸된 경우 이러한 에러가 발생합니다.')
				sleep(1)

	elif g_Tag_type =='video':
		print('video통과')
		#디시콘인지 확인
		g_img_tag = ''
		try:
			g_img_tag = g_temp_tag['class']
		except:
			try:
				g_docximage_info = waitFileDownload(driver,g_img_tag)
				print(g_docximage_info['size'])
				g_docximage_info['size'] = resizeImage(g_docximage_info['size'])
				print(g_docximage_info)
				galldocx.add_picture(g_docximage_info['file_path'],width= docx.shared.Pt(g_docximage_info['size'][0]), height= docx.shared.Pt(g_docximage_info['size'][1]))
				paragraph = galldocx.add_paragraph()
				os.remove(g_docximage_info['file_path'])
			except:
				print('###############이미지를 불러오는 중에 에러가 발생했습니다.')
				sleep(1)

		else:
			print("@@@@@@@@@@@@@디시콘 다운중@@@@@@@@@@@")
			try:
				print(g_temp_tag['data-src'])
			except:
				pass
			else:
				g_img_tag = g_temp_tag['data-src']
				#새탭 열기
				try:
					g_docximage_info = waitFileDownload(driver,g_img_tag)
					g_docximage_info['size'] = resizeImage(g_docximage_info['size'])
					galldocx.add_picture(g_docximage_info['file_path'],width= docx.shared.Pt(g_docximage_info['size'][0]), height= docx.shared.Pt(g_docximage_info['size'][1]))
					paragraph = galldocx.add_paragraph()

					os.remove(g_docximage_info['file_path'])
				except:
					print('###############디시콘을 불러오는 중에 에러가 발생했습니다.')
					print('###############디시콘이 소멸된 경우 이러한 에러가 발생합니다.')
					sleep(1)



		#속성 확인 해야함!
		#속성 확인 종료
		for g_tag_in_tag in list(g_temp_tag):
			listingTag(g_paragraph,g_tag_in_tag,g_TEXT_STYLE)

	#del(g_g_TEXT_STYLE)#안전? 하게 제거

#셀레니움 실행
chrome_options = webdriver.ChromeOptions() # 크롬 브라우저 환경설정
settings = {
       "recentDestinations": [{
            "id": "Save as PDF",
            "origin": "local",
            "account": "",
        }],
        "selectedDestinationId": "Save as PDF",
        "version": 2
    }
prefs = {'printing.print_preview_sticky_settings.appState': json.dumps(settings)}
chrome_options.add_experimental_option('prefs', prefs)
#다운로드 경로 변경
chrome_options.add_experimental_option('prefs', {"download.default_directory": IMG_TEMP_PATH})
chrome_options.add_argument('--kiosk-printing')
driver = webdriver.Chrome(chrome_options=chrome_options,executable_path=DRIVER_PATH)
driver.implicitly_wait(wait_Time)
driver.set_window_size(300, 100) 


ERROR_GALLLIST = []

##############여기서부터 갤 접속 시작########################
def gallURLtoDOCX(g_MAIN_URL):

	global galldocx, docx, ERROR_GALLLIST
	global paragraph, run, driver, chrome_options, settings, prefs, IMG_TEMP_PATH, wait_Time

	#크로미움이 갑자기 닫히는 경우가 발생함
	try:
		_=len(driver.window_handles)
	except:
		print("예상치 못하게 크로미움이 종료되었습니다.. 다시 여는중")
		driver = webdriver.Chrome(chrome_options=chrome_options,executable_path=DRIVER_PATH)
		driver.implicitly_wait(wait_Time)
		driver.set_window_size(300, 100) 
	driver.switch_to.window(driver.window_handles[0])

	alert_error_check = False
	while not alert_error_check:
		try:
			driver.get(g_MAIN_URL)
			html = driver.page_source
		except:
			print("@에러발생: 디시에서 페이지를 보여주지 않습니다. 다시 시도합니다.")
			sleep(2)
		else:
			alert_error_check = True
				
	soup = BeautifulSoup(html,"html.parser")


	print("생성 테스트 완료")

	galldocx = Document()

	#본문 내용 제목으로 저장
	#title_subject저장 후 제목 저장
	title_subject = soup.find(attrs={'class':'title_subject'})
		
	paragraph = galldocx.add_paragraph() 
	paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER #글 중심을 중앙으로 바꾸기
	#해당 글이 삭제되었는지 확인
	try:
		run = paragraph.add_run(title_subject.get_text())
	except:
		print("해당 글이 삭제되었습니다.")
		return None
	
	run.font.name = "Arial"
	run.element.rPr.rFonts.set(qn('w:eastAsia'), '돋움체') #한글 글꼴 바꾸기
	run.font.size = docx.shared.Pt(22)


	paragraph = galldocx.add_paragraph()

	#글 작성자에 대한 정보 추출
	id_info_All = soup.find(attrs={'class':'gallview_head clear ub-content'})
	writer_id_info = getIdInfo(id_info_All)
	id_info_output = "작성시간:"+writer_id_info['date']+', 작성자:'+writer_id_info['user-nickname']+'('+writer_id_info['user']+')'

	run = paragraph.add_run(id_info_output)
	run.font.name = "Arial" #영어 글꼴 바꾸기
	run.element.rPr.rFonts.set(qn('w:eastAsia'), '돋움체') #한글 글꼴 바꾸기
	run.font.size = docx.shared.Pt(10)
	run.bold = True
	run = paragraph.add_run('\n원본 링크: '+g_MAIN_URL)
	setText(run)



	print('본문 작성중..')



	#글 본문을 docx로 바꾸어주는 코드입니다
	#테스트 목적으로 써서 좀 난해합니다..
	write_div = soup.find(attrs={'class':'writing_view_box'})
	image_in_comment = False
	for temp_tag in list(write_div):
		print(list(temp_tag), len(temp_tag),temp_tag.name)
		#print(temp_tag.name)
		try:
			if str(temp_tag.name) != 'script':
				if temp_tag['id'] == 'zzbang_div' or temp_tag['id'] == 'zzbang_ad':
					print('광고 인식 완료. 긁어오지 않습니다')
			else:
				print('스크립트 인식 완료. 긁어오지 않습니다')
		except:
			try:	
				for tag_in_tag in list(temp_tag):
					print(list(tag_in_tag),tag_in_tag.name,tag_in_tag.string)
				print()
				print(temp_tag.get_text())

			except:
				print("에러발생, 다음으로 넘어갑니다.")
			else:
				for tag_in_tag in list(temp_tag):
					paragraph = galldocx.add_paragraph()
					listingTag(paragraph,tag_in_tag)
			
	#존재하는 태그: img, span, font, br, b, i, strike, a

	print('댓글 작성중..')
	paragraph = galldocx.add_paragraph()
	run = paragraph.add_run('\n\n※댓글리스트(마지막 페이지만 표시됩니다.)※\n')
	setText(run,textstyle=DEFAULT_COMMENT_TEXT_STYLE)
	try:
		comment_list = soup.find(attrs={'class':'cmt_list'})
		comment_list = list(comment_list)
	except:
		comment_list = []
		print("@@댓글이 없거나, 댓글을 불러오는 과정에서 문제가 발생했습니다.")

	real_comment_list = []
	image_in_comment = True
	for comment_html in comment_list:
		comment_type_i = ''

		#댓글 답글 댓글돌이 분류
		try:
			comment_type_i = comment_html['class']
		except:
			comment_html = comment_html.find(attrs={'class':'reply_list'})
			for reply_i in list(comment_html):
				#삭제된 댓글인지 확인
				try:
					comment_type_i = len(comment_html.find(attrs={'class':'del_reply'}))
				except:
					real_comment_list.append(('reply',reply_i))
					print('답글입니다.')
				else:
					real_comment_list.append(('del_reply',' '))
					print('삭제된 답글입니다.')
		else:
			if 'ub-content' in comment_type_i and 'dory' not in comment_type_i:
				
				#삭제된 댓글인지 확인
				try:
					comment_type_i = len(comment_html.find(attrs={'class':'del_reply'}))
				except:
					real_comment_list.append(('ub-content',comment_html))
					print('댓글입니다.')
				else:
					real_comment_list.append(('del_ub-content',' '))
					print('삭제된 댓글입니다.')
			else:
				print("댓글로 아닌걸로 인식함",comment_type_i)


	for comment_type, comment_html in real_comment_list:
		run = paragraph.add_run('')
		if 'del' not in comment_type:

			comment_info = getIdInfo(comment_html,g_is_comment=True)
			if comment_type == 'reply':
				run = paragraph.add_run('\t┗답글 : ')
				setText(run,textstyle=DEFAULT_COMMENT_TEXT_STYLE)
			else:
				run = paragraph.add_run('댓글 : ')
				setText(run,textstyle=DEFAULT_COMMENT_TEXT_STYLE)
			run = paragraph.add_run(comment_info['user-nickname']+'('+comment_info['user']+') \\'+comment_info['date']+' :: ')
			setText(run,textstyle=DEFAULT_COMMENT_TEXT_STYLE)


			if not comment_html.find(attrs={'class':'usertxt ub-word'}):
				comment_html = comment_html.find(attrs={'class':'coment_dccon_img'})
				try:
					_ = comment_html.name
				except:
					run = paragraph.add_run('\t보이스 리플 입니다.')
					setText(run,textstyle=DEFAULT_COMMENT_TEXT_STYLE)
				else:
					listingTag(paragraph,comment_html)
					paragraph = galldocx.add_paragraph()
			else:
				comment_html = comment_html.find(attrs={'class':'usertxt ub-word'})
				listingTag(paragraph,comment_html)
		else:
			if 'reply' in comment_type:
				run = paragraph.add_run('\t┗답글 : 삭제된 답글입니다.')
				setText(run,textstyle=DEFAULT_COMMENT_TEXT_STYLE)
			else:
				run = paragraph.add_run('댓글 : 삭제된 댓글입니다.')
				setText(run,textstyle=DEFAULT_COMMENT_TEXT_STYLE)


		run = paragraph.add_run('\n')
		setText(run)


	paragraph = galldocx.add_paragraph()

	#docx이름에 넣을 갤 정보, 추출
	gall_num = soup.find(attrs={'class':'sch_alliance_box clear'})
	gall_num = gall_num['id'][gall_num['id'].rfind('_')+1:]


	#이름 정해야 함!
	docx_save = False
	title_subject_str = title_subject.get_text()
	while not docx_save:
		try:
			galldocx.save(DOCX_SAVE_PATH + '\\' + changeFileName('['+writer_id_info['date'][2:10]+'] '+title_subject_str+' '+writer_id_info['user-nickname']+'('+writer_id_info['user']+') #'+gall_num+'.docx'))
			
		except:
			if len(title_subject_str) >= 4:
				print('이름이 너무 긴것 같습니다.. 글 제목을 줄입니다. 현재 글제목 길이:',len(title_subject_str))
				title_subject_str = title_subject_str[:-4]+'...'
			else:
				print('파일 이름(경로명 포함)이 너무 길거나 다른 이유로 종료 되었습니다.')
				print('에러 기록후 넘어갑니다.')
				docx_save=True
				ERROR_GALLLIST.append(title_subject.get_text()+' #'+gall_num)

		else:
			print('해당 글이 docx파일로 정상적으로 저장되었습니다.')
			docx_save = True

########################################################

def getDocxList(path,extension='.docx',getNumber=True):
	result = []
	for tmpName in os.listdir(path):
		tmpNamePath = path + '//' + tmpName
		if os.path.isdir(tmpNamePath):
			result.extend(getDocxList(tmpNamePath))
		elif extension in tmpNamePath[-6:]:
			if getNumber:
				result.append(int(tmpNamePath[tmpNamePath.rindex('#')+1:tmpNamePath.rindex(extension)]))
			else:
				result.append(tmpNamePath)
	return result

already_gall_list = []
for dir_path in file_path_list:
	print("폴더에 있는 파일들 불러오는 중 :",dir_path)
	already_gall_list.extend(getDocxList(dir_path))
	print("불러오기 완료, 현재 파일 수 : ",len(already_gall_list))
print("-모든 파일을 리스트로 불러오는데 성공했습니다.")


best_Page = int(input("긁어올 개념글 페이지 번호(50개 기준..):"))
best_Page_Range = int(input("긁어올 개념글 페이지 개수: "))

already_download = False

gallList = []
for best_Page_current in range(best_Page,best_Page+best_Page_Range):


	alert_error_check = False
	while not alert_error_check:
		try:
			driver.get("https://gall.dcinside.com/board/lists/?id=jumper&page={0}&exception_mode=recommend&list_num={1}".format(best_Page_current,50)) # 수정요망!!
			html = driver.page_source
		except:
			print("@에러발생: 디시에서 페이지를 보여주지 않습니다. 다시 시도합니다.")
			sleep(2)
		else:
			alert_error_check = True

	soup = BeautifulSoup(html,"html.parser")
	gallList.extend(soup.find_all(attrs={'class':'gall_num'}))
	sleep(0.5)

print(gallList)
print("개념글 리스트 수집 성공, 이제 하나씩 긁는 작업 시작..")


for GALL_URL in gallList:
	print("@@@@@갤러리 글 넘버: ",str(GALL_URL.string))
	try:
		GALL_URL = int(GALL_URL.string)
	except:
		print('정상적인 넘버가 아닙니다.')
	else:
		already_download = False
		if int(GALL_URL) in already_gall_list:
			already_download = True
		if not already_download:
			gallURLtoDOCX("https://gall.dcinside.com/board/view/?id=jumper&no={0}".format(GALL_URL)) # 수정요망!!!
		else:
			print("#####이미 있는 파일입니다.",GALL_URL)

#셀레니움 종료
driver.close()
driver.quit()
print('docx에 저장 및 크롬 종료 완료')

print("---에러가 발생해 저장을 못한 글 넘버")
for error_gall_name in ERROR_GALLLIST:
	print(error_gall_name)
print("----------------------------------")
print("현재 긁어온 념글 페이지(50개 기준) : ",best_Page+best_Page_Range-1)
_ = input("엔터키를 누르면 종료됩니다...")
